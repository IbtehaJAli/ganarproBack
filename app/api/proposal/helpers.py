import docx
from htmldocx import HtmlToDocx


# sections = doc.sections
# for section in sections:
#     for paragraph in section.header.paragraphs:
#         print(f"SECTIOn CHECK {paragraph.text}")

def create_proposal_template(data):
    template_page_2 = 'template_version5_page2'
    if data['project_type']['slug'] == "change_order":
        template_page_2 = "template_version5_page2_change_order"
    elif data['project_type']['slug'] == "pressurewashing":
        template_page_2 = "template_version5_page2_pressurewashing"
    files = ['template_page1v2', template_page_2]
    for file in files:
        process_merge_templates(file, data)
    combine_files = [f"project_type_templates/project_type_{data['project_type']['slug']}.docx",
                     f'{template_page_2}_download.docx']
    combine_word_documents(combine_files)


def transform_data(data):
    try:
        bid_amount = data['bid_amount']
        if int(bid_amount) >= 1000:
            bid_amount = f"{int(bid_amount):,}"

    except Exception as e:
        bid_amount = ''

    dic = {
        '{{COMPANY_NAME}}': data['company_name'] or '',
        '{{COMPANY_STREET}}': data['company_street'] or '',
        '{{COMPANY_CITY}}': data['company_city'] or '',
        '{{COMPANY_STATE}}': data['company_state'] or '',
        '{{COMPANY_ZIP}}': data['company_zip'] or '',
        '{{COMPANY_STATE_SHORT}}': data['company_state_short'] or '',
        '{{COMPANY_CONTACT_EMAIL}}': data['company_contact_email'] or '',
        '{{PROJECT_NAME}}': data['project_name'] or '',
        '{{PROJECT_STREET}}': data['project_street'] or '',
        '{{PROJECT_CITY}}': data['project_city'] or '',
        '{{PROJECT_STATE}}': data['project_state'] or '',
        '{{PROJECT_ZIP}}': data['project_zip'] or '',
        '{{PROPOSAL_POINT_CONTACT_NAME}}': data['proposal_point_contact_name'] or '',
        '{{PROPOSAL_POINT_CONTACT_PHONE}}': data['proposal_point_contact_phone'] or '',
        '{{PROPOSAL_POINT_CONTACT_EMAIL}}': data['proposal_point_contact_email'] or '',
        '{{CURRENT_DATE}}': data['current_date'] or '',
        '{{CUSTOMER_COMPANY_NAME}}': data['customer_company_name'] or '',
        '{{CUSTOMER_STREET}}': data['customer_street'] or '',
        '{{CUSTOMER_CITY}}': data['customer_city'] or '',
        '{{CUSTOMER_STATE}}': data['customer_state'] or '',
        '{{CUSTOMER_ZIP}}': data['customer_zip'] or '',
        '{{PROJECT_CONTACT_1_NAME}}': data['project_contact_1_name'] or '',
        '{{PROJECT_CONTACT_1_PHONE}}': data['project_contact_1_phone'] or '',
        '{{PROJECT_CONTACT_1_EMAIL}}': data['project_contact_1_email'] or '',
        '{{PROJECT_CONTACT_2_NAME}}': data['project_contact_2_name'] or '',
        '{{PROJECT_CONTACT_2_PHONE}}': data['project_contact_2_phone'] or '',
        '{{PROJECT_CONTACT_2_EMAIL}}': data['project_contact_2_email'] or '',
        '{{JOB_SITE_CONTACT_NAME}}': data['job_site_contact_name'] or '',
        '{{JOB_SITE_CONTACT_PHONE}}': data['job_site_contact_phone'] or '',
        '{{JOB_SITE_CONTACT_EMAIL}}': data['job_site_contact_email'] or '',
        '{{BID_AMOUNT}}': bid_amount,
        '{{PROJECT_TYPE_TEMPLATE}}': data['project_type']['name'],
    }

    return dic


def combine_word_documents(files):
    merged_document = docx.Document('documents/template_page1v2_download.docx')

    for index, file in enumerate(files):
        sub_doc = docx.Document(f"documents/{file}")

        # Don't add a page break if you've reached the last file.
        if index < len(files) - 1:
            pass
            # sub_doc.add_page_break()

        for element in sub_doc.element.body:
            merged_document.element.body.append(element)

    merged_document.save('documents/mergedv1.docx')


def process_merge_templates(file_path, data):
    doc = docx.Document(f"documents/{file_path}.docx")
    dic = transform_data(data)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for key in dic.keys():
                        text = p.text
                        if key in text:
                            text = text.replace(key, dic[key])
                            p.text = text

        for p in doc.paragraphs:
            text = p.text
            for key in dic.keys():
                if key in text:
                    text = text.replace(key, dic[key])
                    p.text = text

    sections = doc.sections
    for section in sections:
        for paragraph in section.header.paragraphs:
            text = paragraph.text
            for key in dic.keys():
                if key in text:
                    text = text.replace(key, dic[key])
                    paragraph.text = text
    sections = doc.sections
    for section in sections:
        for paragraph in section.footer.paragraphs:
            text = paragraph.text
            for key in dic.keys():
                if key in text:
                    text = text.replace(key, dic[key])
                    paragraph.text = text

    doc.save(f"documents/{file_path}_download.docx")
